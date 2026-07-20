from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()
style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11.5)
BRAND = RGBColor(0x2C,0x3E,0xB8); GREY = RGBColor(0x60,0x66,0x74)

def title(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=BRAND
def subtitle(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=True; r.font.size=Pt(12); r.font.color.rgb=GREY
def note(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=True; r.font.size=Pt(10.5); r.font.color.rgb=GREY
def h1(t):
    doc.add_paragraph(); p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=BRAND
def sub(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=GREY
def stage(t):
    p=doc.add_paragraph(); r=p.add_run('[ '+t+' ]'); r.italic=True; r.font.color.rgb=GREY; r.font.size=Pt(10.5)
def say(runs):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8)
    for text,bold in runs:
        r=p.add_run(text); r.bold=bold; r.font.size=Pt(12)
def bullet(text, lead=None):
    p=doc.add_paragraph(style='List Bullet')
    if lead:
        r=p.add_run(lead+': '); r.bold=True; r.font.size=Pt(11.5)
    r2=p.add_run(text); r2.font.size=Pt(11.5)
def qa(q,a):
    p=doc.add_paragraph(); r=p.add_run('Q: '+q); r.bold=True; r.font.size=Pt(11.5)
    p2=doc.add_paragraph(); p2.paragraph_format.space_after=Pt(8); r2=p2.add_run('You: '+a); r2.font.size=Pt(11.5)
def rule():
    p=doc.add_paragraph(); r=p.add_run('_'*60); r.font.color.rgb=RGBColor(0xD0,0xD4,0xDA)

title('Focal Point Church: Grace Working Session')
subtitle('Demo script plus the questions to answer before you hang up. Read it warm, take your pauses.')
note('Text in [ brackets ] is a stage direction (what to do on screen). Everything else is what you say out loud. Bold is where to lean in. Every number is from Focal Point\'s real Planning Center and Metrics records.')
rule()

h1('Opening: what today is')
stage('Have the Today page open and be logged in before they sit down.')
say([("Thanks for making the time. Before I show you anything, let me tell you what today is, and what it is not. ",False),("This is not a finished product, and it is not a sales pitch.",True),(" You are our founding church partner. I built a working version of Grace on your real data, your people, your Sundays, your numbers, to show you what is possible, and honestly to put it in front of you so we can shape it together.",False)])
say([("So as I walk through it, ",False),("I want your reactions the whole way",True),(": what is right, what is wrong, what is missing, what you would change. Nothing here is locked. Think of this as us co-designing your version. And at the very end, I have a handful of questions I need your help with to turn this from a demo into something that actually runs for you.",False)])
say([("The big idea is simple. This is not a stack of dashboards. It is one assistant, we call her Grace, and she watches your whole church from two directions: ",False),("the front door, guests walking in, and the back door, people quietly drifting out.",True),(" Her job is making sure nobody slips through either one. Let me show you a normal Monday.",False)])

h1('1. Today: the Monday briefing')
stage('You are on the Today page.')
say([("This is the first thing your pastor sees Monday morning, Grace’s handoff, the way a great assistant catches you up before staff meeting.",False)])
stage('Read a couple of lines of the note, or paraphrase the top.')
say([("Look at this line: ",False),("the Media Team is not on the schedule for the next three Sundays, and it serves every single week. Nobody set it up yet.",True),(" That is the kind of thing that blows up Saturday night, and she caught it three weeks early.",False)])
say([("Down here is ‘Your Monday’: everything that needs you this week, pulled into one list, and every row links to the page where you act.",False)])

h1('2. Front Desk and Guests: the front door')
stage('Front Desk and Guests page.')
say([("This is the front door: every first-time guest, and where they are on the journey from visiting to belonging.",False)])
say([("Here is the number I would sit with. ",False),("29 of your 30 most recent guests are still just visitors.",True),(" They came, got a welcome, then nothing moved them forward. That is the leak every church has and cannot see. This board is how you close it.",False)])
stage('Approve a welcome draft inline.')
say([("Grace already wrote a welcome in the pastor’s voice. I approve it, and it goes. ",False),("She drafts, a human always sends.",True)])

h1('3. Care and Drift: the back door')
stage('Care and Drift page.')
say([("This is the back door, and Grace watches ",False),("three doors here",True),(": families whose kids stopped coming, people who stopped serving, and people who went quiet in their small group.",False)])
stage('Open the Leader Digest and tap ‘I reached out’ live.')
say([("Your leaders do not live in a dashboard. Grace emails each leader just their team’s list, and they ",False),("tap one button when they have reached out.",True),(" That tap moves the card on your board. Nobody has to nag anybody.",False)])
stage('Point to the Manuel Luciani card.')
say([("And here is the moment. This man shows up in ",False),("all three doors at once",True),(": stopped serving, quiet in his group, and now missing Sundays. That is not a maybe. That is someone leaving, and Grace caught it early. No spreadsheet connects those three dots.",False)])

h1('4. Sundays and Comms: getting Sunday done')
stage('Sundays and Comms page.')
say([("This is where Sunday gets staffed. Here is the Media Team, ",False),("not scheduled for three weeks running.",True),(" A normal tool stays silent when a team is forgotten. Grace knows it serves every week, so she flags it ahead.",False)])
stage('Point to a short team and its suggestions.')
say([("And when you fill a spot, watch how careful she is. She suggests real people who have actually served that team, skips anyone near burnout, leaves your staff out of it, and even tells you ",False),("which instruments the band is short.",True)])

h1('5. Insights: she sees your church')
stage('Insights page. Your strongest section, slow down.')
say([("First, the big picture. ",False),("You have grown every single year: 602 a weekend in 2022, over 1,000 now. Up 67% in five years, not one down year.",True),(" And this year is up about 6% over the same stretch last year. Your real data, both years.",False)])
say([("Second, the one that matters most. ",False),("629 people raised their hand for salvation last year, and you are already at 249 this year, on pace to pass 700.",True),(" That is the mission, the number behind everything else here.",False)])
say([("Third, the honest read. About 1 in 4 of your core serve. Your ",False),("Brazilian groups are your strongest",True),(", ten a week. And you have 12,000 records in Planning Center but only about 1,750 real members. Grace works off your real core, not the noise.",False)])

h1('The close')
say([("So that is it. One assistant, watching every door, on your real data. Five years of growth, over 600 salvations a year, and nobody slipping out the back unnoticed.",False)])
say([("Everything you saw is running today. The next piece, where the tracking advances on its own, is the first thing we build together. This is founding-client work, I build it hand in hand with you, not hand you a login and walk away.",False)])

h1('What I need from you before we wrap')
say([("Before we hang up, I have a few things I need to make this real. Most of it comes down to one thing: ",False),("everything Grace does routes to a person, so I need to know who.",True),(" Let me walk through it.",False)])
sub('Who owns what (your point people)')
bullet('who is my main go-to at Focal Point, the person I coordinate everything with.', lead='Main point of contact')
bullet('who leads each serving team? Grace routes stopped-serving and burnout flags to them.', lead='Ministry team leaders')
bullet('who leads each growth group? Group-drift flags route to them for the fall.', lead='Group leaders')
bullet('who owns building the Sunday schedule and filling the gaps?', lead='Volunteer scheduling')
bullet('who is the admin that approves and sends the newsletter and recap?', lead='Communications')
bullet('when a family goes quiet and does not respond, who makes the personal call? Pastor Mark, a care pastor, a care team?', lead='Care and follow-up')
bullet('whose voice are the messages written in, and who signs off before anything sends?', lead='Voice and approval')
sub('A few calls only you can make')
bullet('how many missed Sundays before Grace flags a family as drifting? (I used three.)')
bullet('how often is too often to serve, before it is burnout? (I used more than twice a month.)')
bullet('how long should a silent case sit before it escalates to a call?')
bullet('the minimum people each team truly needs to run, so Grace only flags a real shortage.')
sub('Access')
bullet('the full Planning Center access, so giving participation and the full groups picture light up. Christina, is that something you can set up?')
sub('Direction')
bullet('of everything you saw today, what is the one thing that, if it were running this Monday, would change your week? That is what we build first.')
sub('Next step')
bullet('who is my point person, and when do we reconnect to set this up?')

h1('If they ask (answers ready)')
qa("Is all of this really our data?","Yes, every number. Attendance and growth from your Metrics workbooks, guests, roster, serving, and groups from Planning Center. Salvations from your own weekly challenge tracking, 629 last year.")
qa("Why does it say July 5, not this past Sunday?","Your in-person attendance comes from your weekly summary sheet, which runs through July 5. It updates the moment you send the new week. Everything from Planning Center directly is current as of today.")
qa("Are groups connected?","Fully. Membership and monthly attendance are live, and the group-drift is real. We only count the season, not summer, since groups do not meet then.")
qa("Where is giving?","Giving is the one piece still to connect, and even then we only ever show participation, never dollar amounts. That is deliberate. Everything else is live.")
qa("Does Grace send things without us?","Never. She drafts, a person approves and sends. And there are no giving dollar figures anywhere on screen. Both are deliberate.")

rule()
note('Before you start: be logged in (the dashboard is private, real member data). Have the Before Grace / With Grace toggle ready on Today. The three moments most likely to land: the forgotten Media Team, tapping ‘I reached out’ in the leader digest, and Manuel across all three doors.')

out='docs/Focal-Point-Grace-Demo-Script.docx'
doc.save(out); print('saved',out)
